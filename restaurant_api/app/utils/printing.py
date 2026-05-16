import io
from typing import List
from docx import Document
from docx.shared import Pt
from openpyxl import Workbook


def menu_to_docx_bytes(menu, dishes: List[dict]) -> bytes:
    """
    Возвращает .docx: меню с таблицей блюд и ингредиентов.
    """
    doc = Document()
 
    doc.add_heading(f"Menu: {menu.name}", level=1)

   
    doc.add_paragraph(f"Menu ID: {menu.id}")

    table = doc.add_table(rows=1, cols=5)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "№"
    hdr_cells[1].text = "Название"
    hdr_cells[2].text = "Описание"
    hdr_cells[3].text = "Цена"
    hdr_cells[4].text = "Ингредиенты"

   
    for i, d in enumerate(dishes, start=1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(i)
        row_cells[1].text = d.get("name", "")
        row_cells[2].text = d.get("description") or ""
        row_cells[3].text = f"{d.get('price', 0):.2f}"
        row_cells[4].text = ", ".join(d.get("ingredients", []))

   
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.read()


def menu_to_xlsx_bytes(menu, dishes: List[dict]) -> bytes:
    """
    Возвращает .xlsx: таблица со списком блюд.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = f"Menu_{menu.id}"

   
    ws.append([f"Menu: {menu.name}"])
    ws.append([])  # пустая строка

    
    ws.append(["№", "Название", "Описание", "Цена", "Ингредиенты"])

    for i, d in enumerate(dishes, start=1):
        ws.append([
            i,
            d.get("name", ""),
            d.get("description") or "",
            float(d.get("price", 0)),
            ", ".join(d.get("ingredients", []))
        ])

    
    for column_cells in ws.columns:
        length = max(len(str(cell.value)) if cell.value is not None else 0 for cell in column_cells)
        adjusted_width = (length + 2)
        ws.column_dimensions[column_cells[0].column_letter].width = adjusted_width

    bio = io.BytesIO()
    wb.save(bio)
    bio.seek(0)
    return bio.read()